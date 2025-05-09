import streamlit as st
import pandas as pd
import plotly.express as px
#from deep_translator import GoogleTranslator
#from docx import Document
from fpdf import FPDF
from PyPDF2 import PdfReader
import base64, os, datetime, io, random, uuid, time

# === PAGE CONFIG & NAVBAR ===
st.set_page_config(page_title="STARTEAM Global", layout="wide")
st.markdown("""
<style>
  .topnav { display:flex;justify-content:center;gap:50px;
            background:white;padding:16px;font-size:16px;font-weight:500; }
  .topnav a { text-decoration:none;color:#2563eb; }
  .topnav a:hover { text-decoration:underline; }
  .topbar-shadow { box-shadow:0 2px 4px rgba(0,0,0,0.05);
                   position:sticky;top:0;z-index:1000;background:white; }
</style>
<div class="topbar-shadow">
  <div class="topnav">
    <a href="/?page=Home">Home</a>
    <a href="/?page=Manage EQ List">Manage EQ List</a>
    <a href="/?page=Search">Search</a>
    <a href="/?page=FAQ">FAQ</a>
    <a href="/?page=Translation">Translation</a>
  </div>
</div>
""", unsafe_allow_html=True)





# === HOME ===
eq_list = [  # placeholder
    {"eq_id":"EQ001","customer_name":"ABC","stg_pn":"STG-1","factory_pn":"F-1","engineer":"A","status":"Open","date":"2024-12-01"},
    {"eq_id":"EQ002","customer_name":"XYZ","stg_pn":"STG-2","factory_pn":"F-2","engineer":"B","status":"Closed","date":"2024-12-02"},
]
questions = [
    {"eq_id":"EQ001","description":"Why bubbling?","created_at":"2024-12-01"},
    {"eq_id":"EQ002","description":"Thickness?","created_at":"2024-12-02"},
]

def home_page():
    # --- EQ Data ---
    df_eqs = pd.DataFrame(eq_list)
    df_eqs["date"] = pd.to_datetime(df_eqs["date"])

    # --- Question Data ---
    df_questions = pd.DataFrame(questions)
    df_questions["dueDate"] = pd.to_datetime(df_questions["created_at"]).dt.date
    df_questions_display = (
        df_questions[["eq_id", "description", "dueDate"]]
        .rename(columns={
            "eq_id": "EQ ID",
            "description": "Question",
            "dueDate": "Due Date"
        })
    )

    # --- Chart Data ---
    chart_df = df_eqs["status"].value_counts().reset_index()
    chart_df.columns = ["Status", "Count"]

    # --- Title ---
    st.title("EQ Dashboard 工程问题仪表盘")

    # --- Latest EQs Cards ---
    st.subheader("Latest EQs 最近收到的问题")
    latest_eqs = df_eqs.sort_values("date", ascending=False).head(3)
    cols = st.columns(3)
    for idx, row in latest_eqs.iterrows():
        with cols[idx]:
            st.markdown(f"""
                <div style="background:white;border-radius:10px;padding:1rem;
                            box-shadow:0 1px 4px rgba(0,0,0,0.08);margin-bottom:1rem;
                            font-size:14px;">
                    <div style="font-weight:600;font-size:16px;margin-bottom:6px;">
                        🛠️ EQ ID: {row["eq_id"]}
                    </div>
                    <div><strong>Customer:</strong> {row["customer_name"]}</div>
                    <div><strong>STG P/N:</strong> {row["stg_pn"]}</div>
                    <div><strong>Factory:</strong> {row["factory_pn"]}</div>
                    <div><strong>Engineer:</strong> {row["engineer"]}</div>
                    <div><strong>Status:</strong> {row["status"]}</div>
                    <div style="color:#9ca3af;margin-top:6px;">
                        Received: {row["date"].strftime('%Y-%m-%d')}
                    </div>
                </div>
            """, unsafe_allow_html=True)

    # --- Charts and Recent Questions ---
    col1, col2 = st.columns(2)

    with col1:
        st.subheader("EQs Status Tracking 问题状态追踪")
        fig = px.pie(
            chart_df,
            names="Status",
            values="Count",
            hole=0.6,
            color_discrete_sequence=["#f97316", "#eab308", "#1e3a8a"]
        )
        fig.update_traces(textinfo="label+percent", textposition="inside")
        fig.update_layout(showlegend=True)
        st.plotly_chart(fig, use_container_width=True)

    with col2:
        st.subheader("Recent Questions 最新问题展示")
        st.dataframe(df_questions_display.head(5), use_container_width=True)




# === MANAGE EQ LIST ===
import datetime
import io
import base64
import random
import uuid
import time

# ----- 模擬數據庫 -----

# 模擬顧客數據
customers = [
    {"id": "C001", "name": "ABC Electronics", "contact": "John Smith", "email": "john@abc.com", "phone": "123-456-7890"},
    {"id": "C002", "name": "XYZ Corp", "contact": "Jane Doe", "email": "jane@xyz.com", "phone": "234-567-8901"},
    {"id": "C003", "name": "123 Industries", "contact": "Bob Johnson", "email": "bob@123ind.com", "phone": "345-678-9012"},
    {"id": "C004", "name": "Tech Solutions", "contact": "Alice Brown", "email": "alice@techsol.com", "phone": "456-789-0123"},
    {"id": "C005", "name": "Global Manufacturing", "contact": "Charlie Davis", "email": "charlie@global.com", "phone": "567-890-1234"},
]

# 模擬EQ數據
if 'eq_list' not in st.session_state:
    st.session_state.eq_list = [
        {
            "eq_id": "EQ001",
            "customer_name": "ABC Electronics",
            "customer_pn": "CA-001",
            "stg_pn": "STG-001",
            "factory_pn": "F-001",
            "engineer": "Jack Luo",
            "date": datetime.datetime(2025, 4, 25),
            "via_type": "绿油全塞：IPC 4761 Type VI-b,Plug with solder mask",
            "panel_size": "100*150",
            "base_material": "FR4",
            "solder_mask": "Green",
            "status": "Reviewing",
            "questions": []
        },
        {
            "eq_id": "EQ002",
            "customer_name": "XYZ Corp",
            "customer_pn": "CB-002",
            "stg_pn": "STG-002",
            "factory_pn": "F-002",
            "engineer": "Jack Luo",
            "date": datetime.datetime(2025, 4, 20),
            "via_type": "绿油半塞+全塞：IPC 4761 Type VI-a+b,Plug with solder mask",
            "panel_size": "120*180",
            "base_material": "FR4",
            "solder_mask": "Blue",
            "status": "Pending",
            "questions": []
        },
        {
            "eq_id": "EQ003",
            "customer_name": "123 Industries",
            "customer_pn": "CC-003",
            "stg_pn": "STG-003",
            "factory_pn": "F-003",
            "engineer": "Jack Luo",
            "date": datetime.datetime(2025, 4, 15),
            "via_type": "绿油全塞：IPC 4761 Type VI-b,Plug with solder mask",
            "panel_size": "150*200",
            "base_material": "CEM-3",
            "solder_mask": "Red",
            "status": "Closed",
            "questions": []
        }
    ]

# 模擬問題數據
if 'questions' not in st.session_state:
    st.session_state.questions = [
        {
            "id": "Q001",
            "eq_id": "EQ001",
            "description": "Please confirm if we can use FR4 as the PCB material?",
            "standardized_text": "Please confirm if we can use FR4 as the PCB material? This material will affect the product's heat resistance and stability.",
            "image_path": "https://via.placeholder.com/300x200.png?text=PCB+Sample+1",
            "engineer_suggestion": "Confirm to use FR4.",
            "customer_response": "",
            "status": "Reviewing",
            "created_at": datetime.datetime(2025, 4, 25, 10, 30)
        },
        {
            "id": "Q002",
            "eq_id": "EQ002",
            "description": "Please confirm the pad size.",
            "standardized_text": "Please confirm the pad size meets BGA requirements, currently designed as 0.4mm.",
            "image_path": "https://via.placeholder.com/300x200.png?text=PCB+Sample+2",
            "engineer_suggestion": "According to BGA package requirements, the pad size should be adjusted to 0.5mm to ensure good soldering effect.",
            "customer_response": "Agree to adjust to 0.5mm.",
            "status": "Closed",
            "created_at": datetime.datetime(2025, 4, 20, 14, 45)
        },
        {
            "id": "Q003",
            "eq_id": "EQ002",
            "description": "Color selection for Solder Mask .",
            "standardized_text": "Please confirm if the solder mask color for this batch of PCB can be green",
            "image_path": "https://via.placeholder.com/300x200.png?text=PCB+Sample+3",
            "engineer_suggestion": "Suggest to use green solder mask, which is beneficial for visual inspection and repair.",
            "customer_response": "confirm to use green solder mask.",
            "status": "Pending",
            "created_at": datetime.datetime(2025, 4, 20, 15, 30)
        }
    ]

# 歷史EQ問題庫
historical_questions = [
    {
        "id": "HQ001",
        "text": "PCB板材料確認，是否可使用FR4材料？ Please confirm if we can use FR4 as the PCB material?",
        "suggestions": "FR4是標準材料，適用於大多數應用場景，耐熱性好。 FR4 is a standard material suitable for most applications with good heat resistance.",
        "similarity": 0.95
    },
    {
        "id": "HQ002",
        "text": "焊盤尺寸確認，BGA封裝要求尺寸為多少？Pad size confirmation, what is the size required for BGA package?",
        "suggestions": "BGA焊盤尺寸建議為0.5mm，過小會影響焊接質量。Suggested pad size for BGA is 0.5mm, too small will affect soldering quality.",
        "similarity": 0.87
    },
    {
        "id": "HQ003",
        "text": "阻焊層顏色選擇，是否有特殊要求？If there are any special requirements for solder mask color?",
        "suggestions": "標準綠色阻焊層最適合目檢和返修，其他顏色可能增加成本。Standard green solder mask is best for visual inspection and repair, other colors may increase cost.",
        "similarity": 0.92
    },
    {
        "id": "HQ004",
        "text": "PCB厚度確認，是否需要特殊處理？Thickness confirmation, is special treatment required?",
        "suggestions": "標準厚度1.6mm適合大多數應用，特殊要求可選擇0.8mm或2.0mm。Standard thickness 1.6mm is suitable for most applications, special requirements can choose 0.8mm or 2.0mm.",
        "similarity": 0.65
    },
    {
        "id": "HQ005",
        "text": "覆銅厚度確認，內層和外層要求是否不同？Copper thickness confirmation, are there different requirements for inner and outer layers?",
        "suggestions": "標準外層1oz，內層0.5oz，高電流應用可增加覆銅厚度。Standard outer layer 1oz, inner layer 0.5oz, high current applications can increase copper thickness.",
        "similarity": 0.72
    }
]

# ----- 輔助函數 -----

def format_date(date):
    """格式化日期顯示"""
    if isinstance(date, datetime.datetime):
        return date.strftime("%Y-%m-%d")
    return date

def create_status_badge(status):
    """創建狀態標籤"""
    status_lower = status.lower()
    return f"""<span class="status-{status_lower}">{status}</span>"""

def generate_mock_image_url():
    """模擬圖片URL生成"""
    image_urls = [
        "https://via.placeholder.com/300x200.png?text=PCB+Sample+1",
        "https://via.placeholder.com/300x200.png?text=PCB+Sample+2",
        "https://via.placeholder.com/300x200.png?text=PCB+Sample+3",
        "https://via.placeholder.com/300x200.png?text=PCB+Sample+4"
    ]
    return random.choice(image_urls)

def generate_excel(eq_data, questions_data):
    """生成EQ問題單Excel"""
    # 創建Excel文件
    output = io.BytesIO()
    
    # 使用with語句確保正確關閉writer
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
       # 創建工作表
        sheet_name = 'EQ Data'
        
        # 基本信息表 - 轉置以垂直顯示
        info_data = {
            "Item": ["EQ List ID", "Customer Name", "STG Customer P/N", "STG P/N", 
                   "Factory P/N", "Engineer", "Date", "Status", "Base Material", "Solder Mask"],
            "Value": [eq_data["eq_id"], eq_data["customer_name"], eq_data["customer_pn"], 
                  eq_data["stg_pn"], eq_data["factory_pn"], eq_data["engineer"], 
                  format_date(eq_data["date"]), eq_data["status"], 
                  eq_data.get("base_material", ""), eq_data.get("solder_mask", "")]
        }
        
        df_info = pd.DataFrame(info_data)
        
        # 先寫入基本信息
        df_info.to_excel(writer, sheet_name=sheet_name, index=False, startrow=0)
        
        # 空行
        blank_row = pd.DataFrame()
        blank_row.to_excel(writer, sheet_name=sheet_name, index=False, startrow=len(df_info) + 2)
        
        # 問題列表標題
        title_df = pd.DataFrame({"EQ Question List": [""]})
        title_df.to_excel(writer, sheet_name=sheet_name, index=False, startrow=len(df_info) + 3)
        
        # 問題列表
        df_questions = pd.DataFrame([
            {
                "NO.": i+1,
                "Question Description": q["standardized_text"],
                "Suggestion or Proposal": q["engineer_suggestion"],
                "Customer's Reply": q["customer_response"],
                "Status": q["status"]
            } for i, q in enumerate(questions_data)
        ])
        
        # 寫入問題列表，從基本信息下方開始
        df_questions.to_excel(writer, sheet_name=sheet_name, index=False, startrow=len(df_info) + 4)
        
        # 獲取workbook和worksheet以進行格式設置
        workbook = writer.book
        worksheet = writer.sheets[sheet_name]
        
        # 嘗試設置列寬 (如果使用openpyxl引擎)
        try:
            for idx, col in enumerate(['A', 'B', 'C', 'D', 'E', 'F']):
                if idx == 0:  # 序號列較窄
                    worksheet.column_dimensions[col].width = 8
                elif idx == 1 or idx == 2:  # 問題和建議列較寬
                    worksheet.column_dimensions[col].width = 40
                else:
                    worksheet.column_dimensions[col].width = 20
        except:
            # 如果設置列寬失敗，忽略錯誤
            pass
    
    # 获取Excel内容 - 必须在with语句结束后调用
    output.seek(0)
    excel_data = output.getvalue()
    
    return excel_data

def simulate_excel_import():
    """模擬Excel導入功能"""
    # 隨機生成客戶回覆
    responses = [
        "Confirm to use matertial FR4. 確認使用FR4材料。",
        "Agree to adjust to 0.5mm. 同意調整為0.5mm。",
        "Confirm to use green solder mask. 確認使用綠色阻焊層。",
        "Follow width in Gerber. 已確認線寬要求。",
        "Confirm the order qty is 100 and expected delivery leadtime is 6 weeks. 確認PCB數量為100片，交期為6週。"
    ]
    
    # 隨機選擇2-3個回覆
    num_responses = random.randint(2, 3)
    selected_responses = random.sample(responses, num_responses)
    
    return selected_responses

def simulate_ocr(image_file):
    """模擬OCR識別和翻譯功能"""
    # 模擬處理時間
    time.sleep(1)
    
    # 模擬OCR識別結果
    ocr_results = [
        "Please confirm if we can use FR4 as the PCB material? This material will affect the product's heat resistance and stability.",
        "Please confirm if the pad size meets BGA requirements, currently designed as 0.4mm.",
        "Please confirm if the solder mask color for this batch of PCB can be green.",
        "Please confirm if the line width and spacing meet the production requirements.",
        "Please confirm the order qty and expected delivery leadtime."
    ]
    
    # 隨機選擇一個結果
    result = random.choice(ocr_results)
    
    return result

def standardize_text(text):
    """模擬文本標準化處理"""
    # 模擬處理時間
    time.sleep(0.5)
    
    # 簡單返回原文本，實際應用中可能有更複雜的處理
    return text

# 在函數外部添加以下函數來處理圖片上傳
def save_uploaded_file(uploaded_file):
    """保存上傳的文件並返回本地路徑"""
    # 創建臨時目錄（如果不存在）
    import os
    if not os.path.exists("temp_uploads"):
        os.makedirs("temp_uploads")
    
    # 保存文件
    file_path = os.path.join("temp_uploads", uploaded_file.name)
    with open(file_path, "wb") as f:
        f.write(uploaded_file.getbuffer())
    
    return file_path

# ----- 數據操作函數 -----

def generate_question_id():
    """生成新問題ID"""
    return f"Q{uuid.uuid4().hex[:6].upper()}"

def generate_eq_id():
    """生成新EQ ID"""
    return f"EQ{uuid.uuid4().hex[:6].upper()}"

def get_eq_list(status=None):
    """根據狀態獲取EQ列表，如果沒有指定狀態則返回所有EQ"""
    if status:
        filtered_eq = [eq for eq in st.session_state.eq_list if eq["status"] == status]
        return filtered_eq
    return st.session_state.eq_list

def get_eq_details(eq_id):
    """根據ID獲取EQ詳情"""
    for eq in st.session_state.eq_list:
        if eq["eq_id"] == eq_id:
            return eq
    return None

def create_eq(eq_data):
    """創建新的EQ並返回其ID"""
    new_eq_id = generate_eq_id()
    eq_data["eq_id"] = new_eq_id
    eq_data["questions"] = []
    st.session_state.eq_list.append(eq_data)
    return new_eq_id

def update_eq(eq_id, eq_data):
    """更新現有EQ的信息"""
    for i, eq in enumerate(st.session_state.eq_list):
        if eq["eq_id"] == eq_id:
            # 保留問題列表
            questions = st.session_state.eq_list[i]["questions"]
            # 更新數據
            st.session_state.eq_list[i] = eq_data
            st.session_state.eq_list[i]["questions"] = questions
            return True
    return False

def get_eq_questions(eq_id):
    """獲取特定EQ的所有問題"""
    eq_questions = [q for q in st.session_state.questions if q["eq_id"] == eq_id]
    
    # 調試信息
    print(f"獲取EQ {eq_id} 的問題:")
    print(f"- 所有問題: {len(st.session_state.questions)}個")
    print(f"- 該EQ問題: {len(eq_questions)}個")
    
    return eq_questions

def add_question_to_eq(eq_id, question_data):
    """添加新問題到指定EQ"""
    # 創建新問題ID
    question_id = generate_question_id()
    
    # 添加問題
    new_question = {
        "id": question_id,
        "eq_id": eq_id,
        "description": question_data["description"],
        "standardized_text": question_data["standardized_text"],
        "image_path": question_data.get("image_path", ""),
        "engineer_suggestion": question_data.get("engineer_suggestion", ""),
        "customer_response": "",
        "status": question_data.get("status", "Reviewing"),
        "created_at": datetime.datetime.now()
    }
    
    # 添加到問題列表
    st.session_state.questions.append(new_question)
    
    # 確保問題被添加到正確的EQ中
    eq_found = False
    for i, eq in enumerate(st.session_state.eq_list):
        if eq["eq_id"] == eq_id:
            eq_found = True
            if "questions" not in eq:
                st.session_state.eq_list[i]["questions"] = []
            st.session_state.eq_list[i]["questions"].append(question_id)
            break
    
    # 調試信息
    print(f"添加問題到EQ {eq_id}: {question_id}")
    print(f"EQ是否找到: {eq_found}")
    print(f"當前EQ有 {len([eq for eq in st.session_state.eq_list if eq['eq_id'] == eq_id][0]['questions'])} 個問題")
    print(f"當前問題總數: {len(st.session_state.questions)}")
    
    return question_id

def delete_question(question_id):
    """刪除一個問題"""
    for i, question in enumerate(st.session_state.questions):
        if question["id"] == question_id:
            st.session_state.questions.pop(i)
            return True
    return False

def update_question(question_id, question_data):
    """更新現有問題的信息"""
    for i, question in enumerate(st.session_state.questions):
        if question["id"] == question_id:
            st.session_state.questions[i].update(question_data)
            return True
    return False

def get_customers():
    """獲取所有客戶信息"""
    return customers

def search_historical_questions(query, limit=5):
    """搜索歷史問題庫中的相似問題"""
    # 簡單模擬搜索，實際應用中可能需要更複雜的相似度計算
    sorted_questions = sorted(historical_questions, key=lambda q: q["similarity"], reverse=True)
    return sorted_questions[:limit]

def update_eq_status(eq_id, status):
    """更新EQ的狀態"""
    for i, eq in enumerate(st.session_state.eq_list):
        if eq["eq_id"] == eq_id:
            st.session_state.eq_list[i]["status"] = status
            return True
    return False

def import_customer_responses(eq_id, responses_data):
    """導入客戶對問題的回覆"""
    # responses_data是一個字典，鍵為問題ID，值為客戶回覆內容
    
    for question_id, response in responses_data.items():
        for i, question in enumerate(st.session_state.questions):
            if question["id"] == question_id and question["eq_id"] == eq_id:
                st.session_state.questions[i]["customer_response"] = response
                # 如果有回覆，狀態變為Closed
                if response:
                    st.session_state.questions[i]["status"] = "Closed"
    
    # 檢查是否所有問題都已Closed，如果是，更新EQ狀態
    eq_questions = [q for q in st.session_state.questions if q["eq_id"] == eq_id]
    all_closed = all(q["status"] == "Closed" for q in eq_questions)
    
    if all_closed and eq_questions:
        update_eq_status(eq_id, "Closed")
    
    return True

# ----- UI組件 -----

def show_eq_list_ui(status=None):
    """顯示EQ列表UI"""
    try:
        # 從數據庫獲取EQ列表
        eq_data = get_eq_list(status)
        
        if not eq_data:
            st.info(f"Failed to find {status} EQs" if status else "Failed to find EQs")
            return
        
        # 準備顯示數據
        display_data = []
        for eq in eq_data:
            display_data.append({
                "Select": False,
                "EQ List ID": eq["eq_id"],
                "Customer": eq["customer_name"],
                "STG Customer P/N": eq["customer_pn"],
                "Created Date": format_date(eq["date"]),
                "Status": eq["status"],
                "Questions": len(get_eq_questions(eq["eq_id"]))
            })
        
        # 創建DataFrame
        df = pd.DataFrame(display_data)
        
        # 顯示數據表格
        edited_df = st.data_editor(df, hide_index=True, 
                                  column_config={
                                      "Select": st.column_config.CheckboxColumn("Select", default=False),
                                      "EQ List ID": st.column_config.TextColumn("EQ ID"),
                                      "Customer": st.column_config.TextColumn("Customer"),
                                      "STG Customer P/N": st.column_config.TextColumn("Customer P/N"),
                                      "Created Date": st.column_config.TextColumn("Created Date"),
                                      "Status": st.column_config.TextColumn("Status"),
                                      "Questions": st.column_config.NumberColumn("Questions")
                                  },
                                  disabled=["EQ List ID", "Customer", "STG Customer P/N", "Created Date", "Status", "Questions"])
        
        # 獲取選擇的EQ
        selected_eq_ids = edited_df[edited_df["Select"]]["EQ List ID"].tolist()
        
        if selected_eq_ids:
            if st.button("Edit the EQ List"):
                # 僅取第一個選中的EQ進行編輯
                st.session_state.selected_eq_id = selected_eq_ids[0]
                st.session_state.current_view = "edit_eq"
                st.rerun()
                
        return eq_data
    
    except Exception as e:
        st.error(f"Error to load EQ: {str(e)}")
        return []

def show_eq_details_ui(eq_id=None):
    """顯示EQ詳細信息UI"""
    if not eq_id and "selected_eq_id" in st.session_state:
        eq_id = st.session_state.selected_eq_id
    
    if not eq_id:
        st.info("Please select an EQ to view details")
        return
    
    try:
        # 獲取EQ詳情
        eq_details = get_eq_details(eq_id)
        
        if not eq_details:
            st.warning(f"Failed to Find 找不到EQ {eq_id}")
            return
        
        # 顯示狀態徽章
        st.markdown(
            f"<div class='card-header'><h3>EQ Information</h3>{create_status_badge(eq_details['status'])}</div>", 
            unsafe_allow_html=True
        )
        
        # 創建一個帶有修改狀態的副本，以便在保存時與原始狀態比較
        if "edit_eq_details" not in st.session_state:
            st.session_state.edit_eq_details = dict(eq_details)
        
        # 顯示可編輯的EQ詳情
        col1, col2 = st.columns(2)
        
        with col1:
            customer_name = st.text_input("Customer Name", value=st.session_state.edit_eq_details.get("customer_name", ""), key=f"edit_customer_name_{eq_id}")
            st.session_state.edit_eq_details["customer_name"] = customer_name
            
            customer_pn = st.text_input("STG Customer's P/N", value=st.session_state.edit_eq_details.get("customer_pn", ""), key=f"edit_customer_pn_{eq_id}")
            st.session_state.edit_eq_details["customer_pn"] = customer_pn
            
            stg_pn = st.text_input("STG P/N", value=st.session_state.edit_eq_details.get("stg_pn", ""), key=f"edit_stg_pn_{eq_id}")
            st.session_state.edit_eq_details["stg_pn"] = stg_pn
            
            factory_pn = st.text_input("Factory P/N", value=st.session_state.edit_eq_details.get("factory_pn", ""), key=f"edit_factory_pn_{eq_id}")
            st.session_state.edit_eq_details["factory_pn"] = factory_pn
        
        with col2:
            engineer = st.text_input("Issued by Factory Engineer", value=st.session_state.edit_eq_details.get("engineer", ""), key=f"edit_engineer_{eq_id}")
            st.session_state.edit_eq_details["engineer"] = engineer
            
            date = st.date_input("Date", value=st.session_state.edit_eq_details.get("date", datetime.datetime.now()), key=f"edit_date_{eq_id}")
            st.session_state.edit_eq_details["date"] = date
            
            # 定義所有可能的 via_type 選項
            via_type_options = ["绿油全塞：IPC 4761 Type VI-b,Plug with solder mask", 
                            "绿油半塞+全塞：IPC 4761 Type VI-a+b,Plug with solder mask", 
                            "树脂塞孔+电镀填平：IPC 4761 Type VII, Resin plugged and capped",
                            "树脂塞孔+双面开窗：IPC 4761 Type V. Plug with resin", 
                            "树脂塞孔+单面/双面盖油：IPC 4761 Type VI-a/b,Plug with resin",
                            "表面处理后塞孔：IPC 4761 Type III-a plug after surface",
                            "没有塞孔：No via plugging needed"]

            # 嘗試獲取當前 via_type 在選項列表中的索引，如果不存在則使用默認值0
            current_via_type = st.session_state.edit_eq_details.get("via_type", via_type_options[0])
            try:
                via_type_index = via_type_options.index(current_via_type)
            except ValueError:
                # 如果當前值不在選項列表中，使用默認值
                via_type_index = 0
                # 順便更新 session_state 中的值
                st.session_state.edit_eq_details["via_type"] = via_type_options[0]

            # 使用嚴謹的方式設置選擇框
            via_type = st.selectbox(
                "Via Plugging Type", 
                options=via_type_options,
                index=via_type_index,
                key=f"edit_via_type_{eq_id}"
            )
            st.session_state.edit_eq_details["via_type"] = via_type
            
            panel_size = st.text_input("Panel Size(mm*mm)", value=st.session_state.edit_eq_details.get("panel_size", ""), key=f"edit_panel_size_{eq_id}")
            st.session_state.edit_eq_details["panel_size"] = panel_size
        
        # 其他字段
        col1, col2 = st.columns(2)
        with col1:
            base_material = st.text_input("Base Material", value=st.session_state.edit_eq_details.get("base_material", ""),
                                        key=f"edit_base_material_{eq_id}")
            st.session_state.edit_eq_details["base_material"] = base_material
        
        with col2:
            solder_mask = st.text_input("Solder Mask", value= st.session_state.edit_eq_details.get("solder_mask", ""),
                                      key=f"edit_solder_mask_{eq_id}")
            st.session_state.edit_eq_details["solder_mask"] = solder_mask
        
        # 狀態更新
        status_options = ["Reviewing", "Pending", "Closed"]
        status_index = status_options.index(st.session_state.edit_eq_details["status"]) if st.session_state.edit_eq_details["status"] in status_options else 0
        
        new_status = st.selectbox(
            "Status",
            options=status_options,
            index=status_index,
            key=f"status_select_{eq_id}"
        )
        st.session_state.edit_eq_details["status"] = new_status
        
        # 添加保存按鈕
        col1, col2 = st.columns(2)
        with col1:
            if st.button("Save the Changes", key="save_eq_button"):
                # 檢查是否有變更
                changes_made = False
                for key, value in st.session_state.edit_eq_details.items():
                    if key != "questions" and key in eq_details and eq_details[key] != value:
                        changes_made = True
                        break
                
                if changes_made:
                    # 保存修改
                    if update_eq(eq_id, st.session_state.edit_eq_details):
                        st.success("EQ Info updated successfully")
                        # 清除編輯狀態
                        if "edit_eq_details" in st.session_state:
                            del st.session_state.edit_eq_details
                        # 重新加載頁面
                        st.rerun()
                    else:
                        st.error("Failed to update EQ info")
                else:
                    st.info("No changes dected")
        
        with col2:
            if st.button("Cancel", key="cancel_edit_button"):
                # 取消編輯，恢復到原始狀態
                if "edit_eq_details" in st.session_state:
                    del st.session_state.edit_eq_details
                st.rerun()
        
        # 顯示問題列表 - 這段代碼需要保持在函數內部，而不是縮進到其他代碼塊中
        st.markdown("## EQ Question List")
        
        # 獲取問題清單
        questions = get_eq_questions(eq_id)
        
        if not questions:
            st.info("No Questions Added Yet ")
            st.markdown("### Total 0 questions  共 0 个问题")
        else:
            # 顯示問題數量
            st.markdown(f"### Total {len(questions)} questions  共 {len(questions)} 个问题")
            
            # 顯示問題表格
            view_mode = st.radio("View Model 查看模式", ["Table", "Details"], horizontal=True)
            
            if view_mode == "Table":
                # 表格視圖
                data = []
                for i, q in enumerate(questions):
                    data.append({
                        "Select": False,
                        "NO": i + 1,
                        "Question": q["standardized_text"][:50] + "..." if len(q["standardized_text"]) > 50 else q["standardized_text"],
                        "Eng. Suggestion": q["engineer_suggestion"][:50] + "..." if len(q["engineer_suggestion"]) > 50 else q["engineer_suggestion"],
                        "Customer Response": q["customer_response"][:50] + "..." if q["customer_response"] else "No Reply Yet",
                        "Status": q["status"],
                        "ID": q["id"]  # 用於標識
                    })
                
                df = pd.DataFrame(data)
                
                # 顯示表格
                display_df = df[["Select", "NO", "Question", "Eng. Suggestion", "Customer Response", "Status"]]
                edited_df = st.data_editor(display_df, hide_index=True,
                                         column_config={
                                             "Select": st.column_config.CheckboxColumn("Select", default=False),
                                             "NO": st.column_config.NumberColumn("NO"),
                                             "Question": st.column_config.TextColumn("Question"),
                                             "Eng. Suggestion": st.column_config.TextColumn("Suggestion"),
                                             "Customer Response": st.column_config.TextColumn("Customer's Reply"),
                                             "Status": st.column_config.TextColumn("Status")
                                         },
                                         disabled=["NO", "Question", "Eng. Suggestion", "Customer Response", "Status"])
                
                # 獲取選擇的問題ID
                selected_question_ids = []
                for idx, row in edited_df.iterrows():
                    if row["Select"]:
                        # 獲取對應的問題ID
                        original_idx = df[df["NO"] == row["NO"]].index[0]
                        q_id = df.iloc[original_idx]["ID"]
                        selected_question_ids.append(q_id)
                
                # 如果選中了問題，顯示操作按鈕
                if selected_question_ids:
                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button("Delete the Selected Question"):
                            for q_id in selected_question_ids:
                                if delete_question(q_id):
                                    st.success(f"Delete {q_id} successfully")
                                else:
                                    st.error(f"Failed to delete {q_id}")
                            st.rerun()
            else:
                # 詳細視圖
                for i, q in enumerate(questions):
                    with st.expander(f"Question {i+1}: {q['standardized_text'][:50]}...", expanded=False):
                        cols = st.columns([3, 1])
                        with cols[0]:
                            st.markdown(f"**Complete Question:** {q['standardized_text']}")
                            st.markdown(f"**Suggestions or Proposals:** {q['engineer_suggestion']}")
                            st.markdown(f"**Customer's Reply:** {q['customer_response'] or 'No Reply Yet'}')")
                            st.markdown(f"**Status:** {q['status']}")
                            
                            # 如果有圖片，顯示圖片
                            if q.get('image_path'):
                                st.image(q['image_path'], caption="Attached Image", width=200)
                        
                        with cols[1]:
                            # 添加刪除按鈕
                            if st.button("Delete", key=f"delete_{q['id']}"):
                                if delete_question(q['id']):
                                    st.success(f"Delete Successfully")
                                    st.rerun()
                                else:
                                    st.error("Failed to Delete")
            
            # 按鈕行動區域
            col1, col2, col3 = st.columns(3)
            
            with col1:
                if eq_details["status"] == "Reviewing" and st.button("Add a Question", key=f"add_question_{eq_id}"):
                    st.session_state.current_view = "add_question"
                    st.rerun()
            
            with col2:
                if eq_details["status"] == "Pending" and st.button("Import Customer's Reply", key=f"import_response_{eq_id}"):
                    st.session_state.current_view = "import_responses"
                    st.rerun()
            
            with col3:
                # 避免使用 not not selected_question_ids 這種容易出錯的寫法
                # 而是直接檢查變數是否存在於作用域中
                copy_button_key = f"copy_create_{eq_id}"
                if st.button("Copy & Create Another List", key=copy_button_key):
                    try:
                        # 檢查變數是否存在並有值
                        if 'selected_question_ids' in locals() and selected_question_ids:
                            # 使用已選擇的問題
                            questions_to_copy = selected_question_ids
                        else:
                            # 如果沒有選擇的問題（或變數不存在），使用所有問題
                            questions_to_copy = [q["id"] for q in questions]
                        
                        st.session_state.selected_question_ids = questions_to_copy
                        st.session_state.eq_template = eq_details
                        st.session_state.current_view = "create_new_eq_from_template"
                        st.rerun()
                    except Exception as e:
                        import traceback
                        st.error(f"Error to create new EQ: {str(e)}")
                        st.error(traceback.format_exc())
            
            # 導出Excel
            if st.button("Export Excel"):
                excel_data = generate_excel(eq_details, questions)
                
                # 創建下載鏈接
                b64 = base64.b64encode(excel_data).decode()
                href = f'<a href="data:application/vnd.openxmlformats-officedocument.spreadsheetml.sheet;base64,{b64}" download="{eq_id}_export.xlsx">Click to Download Excel</a>'
                st.markdown(href, unsafe_allow_html=True)
                
                # 如果是Reviewing狀態，導出後更新為Pending
                if eq_details["status"] == "Reviewing":
                    update_eq_status(eq_id, "Pending")
                    st.success("Status of EQ is updated as Pending")
                    st.rerun()
        
        return eq_details
    
    except Exception as e:
        st.error(f"Error of Loading EQ Details: {str(e)}")
        import traceback
        st.error(traceback.format_exc())
        return None

def show_question_form_ui(eq_id):
    """顯示問題創建表單UI"""
    st.markdown("## Add a New Question")
    
    # 初始化狀態
    if "question_images" not in st.session_state:
        st.session_state.question_images = []
    
    if "ocr_result" not in st.session_state:
        st.session_state.ocr_result = ""
    
    if "standardized_text" not in st.session_state:
        st.session_state.standardized_text = ""
    
    if "search_performed" not in st.session_state:
        st.session_state.search_performed = False
    
    if "search_results" not in st.session_state:
        st.session_state.search_results = []
    
    if "temp_question_data" not in st.session_state:
        st.session_state.temp_question_data = {}
    
    # 問題描述輸入
    question_text = st.text_area("Question Description  问题描述", 
                               value=st.session_state.get("question_text", ""),
                               height=100,
                               placeholder="Please input your questions 在此输入您的问题...")
    st.session_state.question_text = question_text  # 保存到session_state
    st.session_state.temp_question_data["description"] = question_text
    
    # 圖片上傳
    st.markdown("## Upload Images  上传图片")
    uploaded_files = st.file_uploader("Drag or Upload an Image 拖拽或上传图片", 
                                   accept_multiple_files=True, 
                                   type=["jpg", "jpeg", "png"])
    
    # 處理上傳圖片
    if uploaded_files:
        # 初始化新圖片列表
        new_images = []
        for file in uploaded_files:
            # 檢查是否已添加過
            if file.name not in [img["name"] for img in st.session_state.question_images]:
                # 模擬生成圖片URL
                image_url = generate_mock_image_url()
                new_images.append({
                    "name": file.name,
                    "url": image_url  # 使用url作為鍵名
                })
        
        # 添加到已有圖片列表
        if new_images:
            st.session_state.question_images.extend(new_images)
    
    # 顯示已上傳圖片
    if st.session_state.question_images:
        st.markdown("### Uploaded Images  已上传的图片")
        cols = st.columns(3)
        for i, img in enumerate(st.session_state.question_images):
            with cols[i % 3]:
                # 使用url顯示圖片
                st.image(img["url"], caption=img["name"], width=150)
                if st.button(f"Remove", key=f"remove_{i}"):
                    st.session_state.question_images.pop(i)
                    st.rerun()
    
    # OCR按鈕
    # 檢查是否有上傳的圖片
    has_images = len(st.session_state.question_images) > 0
    
    if has_images and st.button("OCR - Generate Standardized Text  OCR - 生成标准化文本"):
        with st.spinner("OCR Processing..."):
            # 使用第一張圖片進行OCR
            image_url = st.session_state.question_images[0]["url"]
            # 模擬OCR處理
            ocr_result = simulate_ocr(image_url)
            st.session_state.ocr_result = ocr_result
            # 標準化文本
            st.session_state.standardized_text = standardize_text(ocr_result)
            st.rerun()  # 重新載入頁面以顯示結果
    
    # 顯示OCR結果
    if st.session_state.ocr_result:
        st.markdown("### OCR Result  OCR结果")
        st.info(st.session_state.ocr_result)
        
        st.markdown("### Standardized Text  标准化问题")
        standardized_text = st.text_area("Please edit if necessary 如需修改请编辑", value=st.session_state.standardized_text, height=100)
        st.session_state.standardized_text = standardized_text
        st.session_state.temp_question_data["standardized_text"] = standardized_text
        
        # *** 第1步：必須先搜索歷史問題 ***
        if not st.session_state.search_performed:
            if st.button("Search Historical EQ", key="search_btn"):
                with st.spinner("Searching 搜索中..."):
                    # 搜索歷史問題
                    search_results = search_historical_questions(standardized_text)
                    st.session_state.search_results = search_results
                    st.session_state.search_performed = True
                    st.rerun()  # 重新載入頁面以顯示結果
        else:
            # 已經執行過搜索，顯示搜索結果
            st.markdown("### Search Results 搜索結果")
            
            if not st.session_state.search_results:
                st.info("No similar questions found 没有找到相似的问题")
            else:
                for i, result in enumerate(st.session_state.search_results):
                    with st.expander(f"Question {i+1}: {result['text'][:50]}...", expanded=i==0):
                        st.markdown(f"**Compelete Question:** {result['text']}")
                        st.markdown(f"**Suggestions:** {result['suggestions']}")
                        st.markdown(f"**Similarity:** {result['similarity']:.2f}")
                        
                        col1, col2 = st.columns([3, 1])
                        with col2:
                            if st.button(f"Select this suggestions or proposal", key=f"use_answer_{i}"):
                                # 使用此答案
                                st.session_state.standardized_text = result['text']
                                st.session_state.temp_question_data["standardized_text"] = result['text']
                                st.session_state.temp_question_data["engineer_suggestion"] = result['suggestions']
                                st.rerun()  # 重新載入頁面以顯示更新
            
            # *** 第2步：未找到答案，添加工程建議 ***
            st.markdown("### No Historical EQ Found?")
            st.markdown("Please input your suggestion or proposal：")
            
            engineer_suggestion = st.text_area("Proposal/Suggestions", height=100, 
                                             placeholder="Please input here...",
                                             value=st.session_state.temp_question_data.get("engineer_suggestion", ""))
            st.session_state.temp_question_data["engineer_suggestion"] = engineer_suggestion
            
            # *** 第3步：保存或取消 ***
            col1, col2, col3 = st.columns(3)
            
            with col1:
                # 在 show_question_form_ui 函數中的保存按鈕處理部分
                if st.button("Save the Question", key="save_question_btn"):
                    # 驗證必填欄位
                    if not st.session_state.question_text:
                        st.error("Please input question description")
                    elif not st.session_state.standardized_text:
                        st.error("Please generate or input standardized text")
                    elif not st.session_state.temp_question_data.get("engineer_suggestion"):
                        st.error("Please input your proposal or suggestion")
                    else:
                        # 所有必填欄位都已填寫，保存問題
                        image_path = st.session_state.question_images[0]["url"] if st.session_state.question_images else ""
                        
                        # 保存問題
                        question_data = {
                            "description": st.session_state.question_text,
                            "standardized_text": st.session_state.standardized_text,
                            "image_path": image_path,
                            "engineer_suggestion": st.session_state.temp_question_data.get("engineer_suggestion", ""),
                            "status": "Reviewing"
                        }
                        
                        # 添加問題並獲取問題ID
                        question_id = add_question_to_eq(eq_id, question_data)
                        
                        # 確保這個問題存在於 session_state.questions 中
                        question_exists = False
                        for q in st.session_state.questions:
                            if q["id"] == question_id:
                                question_exists = True
                                break
                        
                        if question_exists:
                            st.success(f"Questions Saved Successfully (ID: {question_id})")
                        else:
                            st.error("Failed to Save, Please Try Again")
                            
                        # 清空表單
                        st.session_state.question_text = ""
                        st.session_state.question_images = []
                        st.session_state.ocr_result = ""
                        st.session_state.standardized_text = ""
                        st.session_state.search_performed = False
                        st.session_state.search_results = []
                        st.session_state.temp_question_data = {}
                        
                        # 添加選項詢問是否繼續添加問題或返回EQ詳情
                        st.session_state.show_after_save_options = True
                        st.rerun()
            
            with col2:
                if st.button("Clear the Form", key="clear_form_btn"):
                    # 清空表單
                    st.session_state.question_text = ""
                    st.session_state.question_images = []
                    st.session_state.ocr_result = ""
                    st.session_state.standardized_text = ""
                    st.session_state.search_performed = False
                    st.session_state.search_results = []
                    st.session_state.temp_question_data = {}
                    st.rerun()
            
            with col3:
                if st.button("Return to EQ Details", key="after_save_back_btn"):
                    st.session_state.show_after_save_options = False
                    # 確保清除任何可能影響問題顯示的臨時狀態
                    if "edit_eq_details" in st.session_state:
                        del st.session_state.edit_eq_details
                    # 設置視圖狀態並強制重新加載
                    st.session_state.current_view = "edit_eq"
                    st.rerun()
    
    # 保存後選項
    if st.session_state.get("show_after_save_options", False):
        st.markdown("---")
        st.markdown("### Question Saved Successfully!")
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("Add Another Questions", key="continue_add_btn"):
                st.session_state.show_after_save_options = False
                st.rerun()
        
        with col2:
            if st.button("Return to EQ Details", key="after_save_back_btn"):
                st.session_state.show_after_save_options = False
                st.session_state.current_view = "edit_eq"
                st.rerun()

def import_customer_responses_ui(eq_id):
    """導入客戶回覆功能UI"""
    eq_details = get_eq_details(eq_id)
    questions = get_eq_questions(eq_id)
    
    st.markdown("## Import Customer's Responses")
    st.markdown(f"### EQ: {eq_id} - {eq_details['customer_name']}")
    
    # 上傳Excel文件
    uploaded_file = st.file_uploader("Upload Excel with Customer's Reply", type=["xlsx"])
    
    if uploaded_file:
        st.success(f"File Uploaded: {uploaded_file.name}")
        
        # 顯示模擬導入按鈕
        if st.button("Import Customer's Reply"):
            with st.spinner("Reply Processing..."):
                # 獲取當前EQ的問題
                if not questions:
                    st.warning("No questions found for this EQ")
                    return
                
                # 模擬導入客戶回覆
                responses = simulate_excel_import()
                
                # 創建問題ID到回覆的映射
                response_map = {}
                for i, q in enumerate(questions):
                    if i < len(responses):
                        response_map[q["id"]] = responses[i]
                
                # 導入回覆
                if import_customer_responses(eq_id, response_map):
                    st.success("Import Customer's Reply Successfully")
                    
                    # 顯示更新後的問題
                    updated_questions = get_eq_questions(eq_id)
                    
                    # 檢查是否所有問題都已關閉
                    all_closed = all(q["status"] == "Closed" for q in updated_questions)
                    
                    if all_closed:
                        st.success("All questions are closed, EQ status updated to Closed")
                    else:
                        st.info("Some questions are still open, EQ status remains as Pending and Please check")
                    
                    # 返回編輯頁面
                    st.session_state.current_view = "edit_eq"
                    st.rerun()
                else:
                    st.error("Failed to import customer's reply")

def create_new_eq_ui():
    """創建新的EQ"""
    st.markdown("## Create a New EQ  /n 创建新的EQ")
    
    # 基本信息
    st.markdown("### Basic EQ Info  基本信息")
    
    col1, col2 = st.columns(2)
    
    with col1:
        customer_options = [c["name"] for c in customers]
        customer_name = st.selectbox("Customer", options=customer_options)
        customer_pn = st.text_input("STG Customer's P/N")
        stg_pn = st.text_input("STG P/N")
        factory_pn = st.text_input("Factory P/N")
    
    with col2:
        engineer = st.text_input("Issued by Factory Engineer")
        date = st.date_input("Date", value=datetime.datetime.now())
        via_type = st.selectbox("Via Plugging Type", options=["绿油全塞：IPC 4761 Type VI-b,Plug with solder mask", "绿油半塞+全塞：IPC 4761 Type VI-a+b,Plug with solder mask", "树脂塞孔+电镀填平：IPC 4761 Type VII, Resin plugged and capped","树脂塞孔+双面开窗：IPC 4761 Type V. Plug with resin", "树脂塞孔+单面/双面盖油：IPC 4761 Type VI-a/b,Plug with resin","表面处理后塞孔：IPC 4761 Type III-a plug after surface","没有塞孔：No via plugging needed"])
        panel_size = st.text_input("Panel Size(mm*mm)")
    
    col1, col2 = st.columns(2)
    with col1:
        base_material = st.selectbox("Base Material", options=["FR4", "CEM-3", "Aluminum", "High-Tg FR4"])
    with col2:
        solder_mask = st.selectbox("Solder Mask", options=["Green", "Blue", "Red", "Black", "White"])
    
    # 創建EQ按鈕
    if st.button("Add EQ Questions"):
        if not customer_name or not customer_pn:
            st.warning("Please input: Customer 和 STG P/N")
        else:
            # 創建EQ數據
            eq_data = {
                "customer_name": customer_name,
                "customer_pn": customer_pn,
                "stg_pn": stg_pn,
                "factory_pn": factory_pn,
                "engineer": engineer,
                "date": date,
                "via_type": via_type,
                "panel_size": panel_size,
                "base_material": base_material,
                "solder_mask": solder_mask,
                "status": "Reviewing"
            }
            
            # 創建EQ
            new_eq_id = create_eq(eq_data)
            
            st.success(f"Create EQ Successfully: {new_eq_id}")
            
            # 設置當前選中的EQ並進入問題創建頁面
            st.session_state.selected_eq_id = new_eq_id
            st.session_state.current_view = "add_question"
            st.rerun()

def create_new_eq_from_template_ui():
    """從模板創建新的EQ"""
    if "eq_template" not in st.session_state or "selected_question_ids" not in st.session_state:
        st.error("Failed to create new EQ: Missing template or selected questions")
        return
    
    template = st.session_state.eq_template
    selected_question_ids = st.session_state.selected_question_ids
    
    st.markdown("## EQ Create a New EQ from Existing EQ")
    
    # 基本信息
    st.markdown("### Basic EQ Info  基本信息")
    
    col1, col2 = st.columns(2)
    
    with col1:
        customer_options = [c["name"] for c in customers]
        default_customer_index = customer_options.index(template["customer_name"]) if template["customer_name"] in customer_options else 0
        customer_name = st.selectbox("Customer", options=customer_options, index=default_customer_index)
        customer_pn = st.text_input("STG Customer's P/N", value=template["customer_pn"])
        stg_pn = st.text_input("STG P/N", value=template["stg_pn"])
        factory_pn = st.text_input("Factory P/N", value=template["factory_pn"])
    
    with col2:
        engineer = st.text_input("Issued by Factory Engineer", value=template["engineer"])
        date = st.date_input("Date", value=datetime.datetime.now())
        via_type_options = ["绿油全塞：IPC 4761 Type VI-b,Plug with solder mask", "绿油半塞+全塞：IPC 4761 Type VI-a+b,Plug with solder mask", "树脂塞孔+电镀填平：IPC 4761 Type VII, Resin plugged and capped","树脂塞孔+双面开窗：IPC 4761 Type V. Plug with resin", "树脂塞孔+单面/双面盖油：IPC 4761 Type VI-a/b,Plug with resin","表面处理后塞孔：IPC 4761 Type III-a plug after surface","没有塞孔：No via plugging needed"]
        default_via_index = via_type_options.index(template["via_type"]) if template["via_type"] in via_type_options else 0
        via_type = st.selectbox("Via Plugging Type", options=via_type_options, index=default_via_index)
        panel_size = st.text_input("Panel Size(mm*mm)", value=template["panel_size"])
    
    col1, col2 = st.columns(2)
    with col1:
        base_material_options = ["FR4", "CEM-3", "Aluminum", "High-Tg FR4"]
        default_base_index = base_material_options.index(template["base_material"]) if template["base_material"] in base_material_options else 0
        base_material = st.selectbox("Base Material", options=base_material_options, index=default_base_index)
    with col2:
        solder_mask_options = ["Green", "Blue", "Red", "Black", "White"]
        default_solder_index = solder_mask_options.index(template["solder_mask"]) if template["solder_mask"] in solder_mask_options else 0
        solder_mask = st.selectbox("Solder Mask", options=solder_mask_options, index=default_solder_index)
    
    # 選中的問題
    st.markdown("### Selected Questions  已选中的问题")
    
    selected_questions = []
    for q_id in selected_question_ids:
        for q in st.session_state.questions:
            if q["id"] == q_id:
                selected_questions.append(q)
                break
    
    if not selected_questions:
        st.warning("No Questions Selected")
    else:
        # 顯示選中的問題
        for i, q in enumerate(selected_questions):
            with st.expander(f"Question {i+1}: {q['standardized_text'][:50]}...", expanded=i==0):
                st.markdown(f"**Complete Question:** {q['standardized_text']}")
                st.markdown(f"**Proposal or Suggestion:** {q['engineer_suggestion']}")
    
    # 創建EQ按鈕
    if st.button("Create New EQ"):
        if not customer_name or not customer_pn:
            st.warning("Please input: Customer 和 STG P/N")
        else:
            # 創建EQ數據
            eq_data = {
                "customer_name": customer_name,
                "customer_pn": customer_pn,
                "stg_pn": stg_pn,
                "factory_pn": factory_pn,
                "engineer": engineer,
                "date": date,
                "via_type": via_type,
                "panel_size": panel_size,
                "base_material": base_material,
                "solder_mask": solder_mask,
                "status": "Reviewing"
            }
            
            # 創建EQ
            new_eq_id = create_eq(eq_data)
            
            # 複製選中的問題到新EQ
            for q in selected_questions:
                add_question_to_eq(new_eq_id, {
                    "description": q["description"],
                    "standardized_text": q["standardized_text"],
                    "image_path": q["image_path"],
                    "engineer_suggestion": q["engineer_suggestion"],
                    "status": "Reviewing"
                })
            
            st.success(f"Successfully Create EQ List: {new_eq_id} and Copy {len(selected_questions)} Questions")
            
            # 清除臨時數據
            st.session_state.pop("eq_template", None)
            st.session_state.pop("selected_question_ids", None)
            
            # 設置當前選中的EQ並進入詳情頁面
            st.session_state.selected_eq_id = new_eq_id
            st.session_state.current_view = "edit_eq"
            st.rerun()


# === create_eq_page PAGE ===
def create_eq_page():
    st.title("Create & Manage EQ Lists")
    if "current_view" not in st.session_state:
        st.session_state.current_view = "main"

    if st.session_state.current_view == "main":
        col1, col2 = st.columns(2)
        if col1.button("Manage Existing EQ Lists", use_container_width=True):
            st.session_state.current_view = "manage_eq"
            st.rerun()
        if col2.button("Create New EQ List", use_container_width=True):
            st.session_state.current_view = "create_new_eq"
            st.rerun()

    elif st.session_state.current_view == "manage_eq":
        if st.button("← Back to Main"):
            st.session_state.current_view = "main"
            st.rerun()
        tabs = st.tabs(["Reviewing","Pending","All"])
        with tabs[0]:
            show_eq_list_ui("Reviewing")
        with tabs[1]:
            show_eq_list_ui("Pending")
        with tabs[2]:
            show_eq_list_ui()

    elif st.session_state.current_view == "edit_eq":
        if st.button("← Back to EQ Lists"):
            st.session_state.current_view = "manage_eq"
            st.rerun()
        show_eq_details_ui()

    elif st.session_state.current_view == "add_question":
        if st.button("← Back to EQ Detail"):
            st.session_state.current_view = "edit_eq"
            st.rerun()
        show_question_form_ui(st.session_state.selected_eq_id)

    elif st.session_state.current_view == "import_responses":
        if st.button("← Back to EQ Detail"):
            st.session_state.current_view = "edit_eq"
            st.rerun()
        import_customer_responses_ui(st.session_state.selected_eq_id)

    elif st.session_state.current_view == "create_new_eq":
        if st.button("← Back to Main"):
            st.session_state.current_view = "main"
            st.rerun()
        create_new_eq_ui()

    elif st.session_state.current_view == "create_new_eq_from_template":
        if st.button("← Back to EQ Detail"):
            st.session_state.current_view = "edit_eq"
            st.rerun()
        create_new_eq_from_template_ui()

# === SEARCH ===
def search_page():
    st.title("EQ List - STARTEAM Global")

    # — Data Setup —
    df = pd.DataFrame(eq_list)
    df["date"] = pd.to_datetime(df["date"])

    df.rename(columns={
        "eq_id": "ID",
        "customer_name": "Customer",
        "customer_pn": "Customer P/N",
        "stg_pn": "P/N",
        "factory_pn": "Factory",
        "engineer": "Engineer Name",
        "date": "Changed",
        "base_material": "Base Material",
        "solder_mask": "Solder Mask",
        "panel_size": "Delivery Panel Size",
        "via_type": "Via Type",
        "status": "EQ Status"
    }, inplace=True)

    # — Pagination State —
    if "current_page" not in st.session_state:
        st.session_state.current_page = 1

    # — Filters —
    f1, f2, f3, f4, f5 = st.columns([2, 1, 1, 1, 1])
    with f1:
        keyword = st.text_input("Search Customer Name, Customer P/N, or Project")
    with f2:
        start_date = st.date_input("Start Date", value=None)
    with f3:
        end_date = st.date_input("End Date", value=None)
    with f4:
        status_filter = st.selectbox("All Status", ["All"] + sorted(df["EQ Status"].dropna().unique()))
    with f5:
        factory_filter = st.selectbox("All Factories", ["All"] + sorted(df["Factory"].dropna().unique()))

    f6, f7, f8, f9 = st.columns([2, 1, 1, 1])
    with f6:
        item_code = st.text_input("Search Item Code/STG P/N")
    with f7:
        engineer_team = st.selectbox("Engineer Team", ["All", "Team A", "Team B", "Team C"])  # placeholder
    with f8:
        engineer_name = st.selectbox("Engineer Name", ["All"] + sorted(df["Engineer Name"].dropna().unique()))
    with f9:
        cs_name = st.selectbox("CS Name", ["All", "Emily Wong", "Chris Liu", "Nina Yeung"])  # placeholder

    # — Apply Filters —
    filtered_df = df.copy()
    if keyword:
        k = keyword.lower()
        filtered_df = filtered_df[
            filtered_df["Customer"].str.lower().str.contains(k, na=False) |
            filtered_df["Customer P/N"].str.lower().str.contains(k, na=False)
        ]
    if item_code:
        filtered_df = filtered_df[
            filtered_df["P/N"].str.lower().str.contains(item_code.lower(), na=False)
        ]
    if status_filter != "All":
        filtered_df = filtered_df[filtered_df["EQ Status"] == status_filter]
    if factory_filter != "All":
        filtered_df = filtered_df[filtered_df["Factory"] == factory_filter]
    if engineer_name != "All":
        filtered_df = filtered_df[filtered_df["Engineer Name"] == engineer_name]
    if cs_name != "All":
        filtered_df = filtered_df[
            filtered_df["Engineer Name"].str.contains(cs_name.split()[0], na=False)
        ]
    if start_date:
        filtered_df = filtered_df[filtered_df["Changed"] >= pd.to_datetime(start_date)]
    if end_date:
        filtered_df = filtered_df[filtered_df["Changed"] <= pd.to_datetime(end_date)]

    # — Pagination Setup —
    bottom_cols = st.columns([10, 1])
    with bottom_cols[1]:
        page_size = st.selectbox("Rows per page", [5, 10, 20], index=1, key="page_size")

    total_pages = max(1, (len(filtered_df) - 1) // page_size + 1)
    if st.session_state.current_page > total_pages:
        st.session_state.current_page = total_pages

    start_idx = (st.session_state.current_page - 1) * page_size
    end_idx = start_idx + page_size
    page_data = filtered_df.iloc[start_idx:end_idx]

    # — Display Table —
    st.dataframe(page_data, use_container_width=True)

    # — Pagination Buttons —
    col_prev, col_page, col_next = st.columns([1, 6, 1])
    with col_prev:
        if st.button("Previous", disabled=st.session_state.current_page == 1):
            st.session_state.current_page -= 1
    with col_page:
        st.markdown(
            f"<div style='text-align:center;padding-top:8px;'>"
            f"Page {st.session_state.current_page} of {total_pages}</div>",
            unsafe_allow_html=True
        )
    with col_next:
        if st.button("Next", disabled=st.session_state.current_page == total_pages):
            st.session_state.current_page += 1

    # — Back to Dashboard Link —
    st.markdown("---")
    left, right = st.columns([9, 1])
    with right:
        st.markdown("""
            <a href="/?page=Home" style="
                display:inline-block;
                padding:10px 18px;
                background-color:#003366;
                color:white;
                text-decoration:none;
                border-radius:8px;
                font-size:13px;
                font-weight:500;
                box-shadow:0 1px 2px rgba(0,0,0,0.15);
            ">Back to Dashboard</a>
        """, unsafe_allow_html=True)








# === FAQ PAGE ===
def faq_page():
    import streamlit as st
    import pandas as pd
    from datetime import date

    # --- STYLE FIXES (optional, if you need them only for FAQ) ---
    st.markdown("""
        <style>
            .smaller-title { font-size: 15px !important; font-weight: 600; margin-bottom: 10px; }
            input, select { height: 42px !important; font-size: 14px !important; }
            .stButton>button {
                height: 36px; font-size: 13px; font-weight: 500;
                border-radius: 6px; border: none; width: 100%;
                background-color: #00143f; color: white;
            }
            .stButton>button:hover { background-color: #001030; }
        </style>
    """, unsafe_allow_html=True)

    # --- MOCK DATA ---
    customers = [
        {"name": "Siedle", "count": 128},
        {"name": "Deltec", "count": 96},
        {"name": "ifm",    "count": 85},
        {"name": "Pilz",   "count": 64},
        {"name": "Others", "count": 52},
    ]

    faqs = [
        {"similarity": 95, "question": "HDI板层压工艺中，如何优化压合温度曲线以提高产品良率？", 
         "date": "2024-01-15", "customer": "Siedle", "status": "Closed",    "stg": "STG-1001"},
        {"similarity": 88, "question": "高频PCB阻抗测试中出现较大偏差，可能的原因及解决方案？",
         "date": "2024-01-15", "customer": "Deltec", "status": "Pending",   "stg": "STG-2001"},
        {"similarity": 82, "question": "多层板钻孔后出现毛刺，如何调整工艺参数？",           
         "date": "2024-01-14", "customer": "ifm",    "status": "Pending",   "stg": "STG-3002"},
        {"similarity": 80, "question": "层压后出现鼓包的原因分析？",                       
         "date": "2024-01-14", "customer": "Others", "status": "Closed",    "stg": "STG-4003"},
        {"similarity": 90, "question": "沉金工艺中出现颜色不均匀的问题，如何优化？",      
         "date": "2024-01-13", "customer": "Pilz",   "status": "Reviewing", "stg": "STG-5004"},
    ]

    faq_df = pd.DataFrame(faqs)
    faq_df["date"] = pd.to_datetime(faq_df["date"])

    # --- LAYOUT ---
    left_col, right_col = st.columns([1, 4], gap="large")

    # LEFT: customer filter
    with left_col:
        st.markdown("<div class='smaller-title'>按客户筛选 / Filter by Customers</div>", unsafe_allow_html=True)
        cust_search = st.text_input("", placeholder="🔍 Type customer name...")
        matches = [c for c in customers if cust_search.lower() in c["name"].lower()]
        choice = st.radio("", [f"{c['name']} ({c['count']})" for c in matches], label_visibility="collapsed")

    # RIGHT: other filters + results
    with right_col:
        cols = st.columns([2,2,2,2,2,1])
        kw   = cols[0].text_input("关键词 / Keyword", placeholder="Keyword")
        start = cols[1].date_input("Start", value=date(2024,1,1))
        end   = cols[2].date_input("End",   value=date(2024,12,31))
        stat  = cols[3].selectbox("状态 / Status",
                                     ["Closed", "Pending", "Reviewing"],
                                     index=0)
        stg   = cols[4].text_input("STG P/N", placeholder="STG P/N")
        search = cols[5].button("Search")

        # apply filters
        df = faq_df.copy()
        if choice:
            name = choice.split(" (")[0]
            df = df[df["customer"] == name]
        if kw:    df = df[df["question"].str.contains(kw, case=False)]
        if stat!="All Statuses": df = df[df["status"] == stat]
        if stg:   df = df[df["stg"].str.contains(stg, case=False)]
        df = df[(df["date"] >= pd.to_datetime(start)) & (df["date"] <= pd.to_datetime(end))]

        st.markdown(f"<div style='margin-top:12px;'><b>{len(df)} results found</b></div>", unsafe_allow_html=True)

        for _, row in df.iterrows():
            icon = "✅" if row.status=="Closed" else ("🔄" if row.status=="Pending" else "⏳")
            st.markdown(f"""
                <div style="
                    border:1px solid #e0e0e0; border-radius:10px; padding:12px;
                    margin-bottom:10px; background:white;
                ">
                  <div style="font-size:13px; color:#888;">
                    相似度 <b>{row.similarity}%</b>
                  </div>
                  <div style="font-weight:bold; font-size:16px; color:#222;">
                    {row.question}
                  </div>
                  <div style="font-size:13px; color:#555;">
                    {row.date.date()} | 客户：{row.customer} | 状态：{row.status} {icon} | STG P/N：{row.stg}
                  </div>
                  <div style="text-align:right; color:#888; font-size:12px;">
                    📂 查看详情
                  </div>
                </div>
            """, unsafe_allow_html=True)

# === TRANSLATION ===
def translation_page():
    import streamlit as st
    import base64
    import os
    from deep_translator import GoogleTranslator
    from docx import Document
    from fpdf import FPDF
    from PyPDF2 import PdfReader
    from io import BytesIO

    # --- CUSTOM STYLING FOR NAVY-BLUE DOWNLOAD BUTTON ---
    st.markdown("""
        <style>
            .custom-download button {
                background-color: #001f3f !important;
                color: white !important;
                border: none;
                padding: 0.5rem 1.25rem;
                border-radius: 6px;
                font-weight: 500;
                font-size: 0.875rem;
                box-shadow: 0 1px 3px rgba(0,0,0,0.1);
            }
            .custom-download button:hover {
                background-color: #00143f !important;
            }
        </style>
    """, unsafe_allow_html=True)

    # --- RESET STATE IF TRIGGERED ---
    if "reset_triggered" not in st.session_state:
        st.session_state.reset_triggered = False
    if st.session_state.reset_triggered:
        for k in ["source_lang", "target_lang", "source_lang_text", "target_lang_text", "translated_text_area"]:
            st.session_state.pop(k, None)
        st.session_state.reset_triggered = False
        st.experimental_rerun()

    # --- TWO-TAB INTERFACE ---
    tab1, tab2 = st.tabs(["File Translation", "Text Translation"])
    lang_options = ["auto", "en", "zh-CN", "fr", "de"]

    # --- FILE TRANSLATION ---
    with tab1:
        st.markdown("<br>", unsafe_allow_html=True)
        c1, c2, c3 = st.columns([5, 0.5, 5])
        with c1:
            src = st.selectbox("", lang_options, key="source_lang", label_visibility="collapsed")
        with c2:
            if st.button("⇄", help="Swap languages"):
                a, b = st.session_state.source_lang, st.session_state.target_lang
                st.session_state.source_lang, st.session_state.target_lang = (
                    b if b!="auto" else "en",
                    a if a!="auto" else "en"
                )
                st.experimental_rerun()
        with c3:
            tgt = st.selectbox("", [l for l in lang_options if l!="auto"],
                               index=1, key="target_lang", label_visibility="collapsed")

        uploaded = st.file_uploader("Upload your document", type=["pdf","docx","txt"])
        orig = ""
        if uploaded:
            ext = uploaded.name.split(".")[-1].lower()
            if ext=="pdf":
                reader = PdfReader(uploaded)
                orig = "\n".join(p.extract_text() for p in reader.pages if p.extract_text())
            elif ext=="docx":
                doc = Document(uploaded)
                orig = "\n".join(p.text for p in doc.paragraphs)
            else:
                orig = uploaded.read().decode("utf-8")

        if orig.strip():
            with st.spinner("Translating…"):
                trans = GoogleTranslator(source=src, target=tgt).translate(orig)

            ca, cb = st.columns(2)
            with ca: st.text_area("Original Text", value=orig, height=250)
            with cb: st.text_area("Translated Text", value=trans, height=250)

            dl1, dl2 = st.columns(2)
            with dl1:
                fmt = st.selectbox("", ["txt","docx","pdf"], key="file_format", label_visibility="collapsed")

                def mk_download(txt, f):
                    if f=="txt":
                        b = base64.b64encode(txt.encode()).decode()
                        return f'<div class="custom-download"><a href="data:text/plain;base64,{b}" download="translated.txt"><button>Download</button></a></div>'
                    if f=="docx":
                        buf = BytesIO()
                        d = Document()
                        for line in txt.split("\n"): d.add_paragraph(line)
                        d.save(buf)
                        b = base64.b64encode(buf.getvalue()).decode()
                        return f'<div class="custom-download"><a href="data:application/octet-stream;base64,{b}" download="translated.docx"><button>Download</button></a></div>'
                    # pdf
                    pdf = FPDF(); pdf.add_page(); pdf.set_font("Arial", size=12)
                    for line in txt.split("\n"): pdf.multi_cell(0,10,line)
                    pdf_bytes = pdf.output(dest="S").encode("latin-1")
                    b = base64.b64encode(pdf_bytes).decode()
                    return f'<div class="custom-download"><a href="data:application/pdf;base64,{b}" download="translated.pdf"><button>Download</button></a></div>'

                st.markdown(mk_download(trans, fmt), unsafe_allow_html=True)

            with dl2:
                st.button("Clear", on_click=lambda: st.session_state.update({"reset_triggered": True}))

    # --- TEXT TRANSLATION ---
    with tab2:
        st.markdown("<br>", unsafe_allow_html=True)
        c1, c2, c3 = st.columns([5,0.5,5]) 
        with c1:
            sl = st.selectbox("", lang_options, key="source_lang_text", label_visibility="collapsed")
        with c2:
            if st.button("⇄", key="swap_text", help="Swap languages"):
                x, y = st.session_state.source_lang_text, st.session_state.target_lang_text
                st.session_state.source_lang_text, st.session_state.target_lang_text = (
                    y if y!="auto" else "en",
                    x if x!="auto" else "en"
                )
                st.experimental_rerun()
        with c3:
            tl = st.selectbox("", [l for l in lang_options if l!="auto"],
                              index=1, key="target_lang_text", label_visibility="collapsed")

        left, right = st.columns(2)
        with left:
            user_input = st.text_area("Original Text", height=250, key="user_input")
        with right:
            st.text_area("Translated Text",
                         value=st.session_state.get("translated_text_area",""),
                         height=250, disabled=True)

        if st.button("Translate Text"):
            if user_input.strip():
                with st.spinner("Translating…"):
                    st.session_state.translated_text_area = GoogleTranslator(
                        source=sl, target=tl
                    ).translate(user_input)
            else:
                st.warning("Please enter some text to translate.")

        if st.session_state.get("translated_text_area"):
            d1, d2 = st.columns(2)
            with d1:
                fmt2 = st.selectbox("", ["txt","docx","pdf"], key="text_format", label_visibility="collapsed")

                def mk_txt_download(txt, f):
                    if f=="txt":
                        b = base64.b64encode(txt.encode()).decode()
                        return f'<div class="custom-download"><a href="data:text/plain;base64,{b}" download="text.txt"><button>Download</button></a></div>'
                    if f=="docx":
                        buf = BytesIO()
                        d = Document()
                        for line in txt.split("\n"): d.add_paragraph(line)
                        d.save(buf)
                        b = base64.b64encode(buf.getvalue()).decode()
                        return f'<div class="custom-download"><a href="data:application/octet-stream;base64,{b}" download="text.docx"><button>Download</button></a></div>'
                    pdf = FPDF(); pdf.add_page(); pdf.set_font("Arial",size=12)
                    for line in txt.split("\n"): pdf.multi_cell(0,10,line)
                    pbytes = pdf.output(dest="S").encode("latin-1")
                    b = base64.b64encode(pbytes).decode()
                    return f'<div class="custom-download"><a href="data:application/pdf;base64,{b}" download="text.pdf"><button>Download</button></a></div>'

                st.markdown(mk_txt_download(st.session_state.translated_text_area, fmt2),
                            unsafe_allow_html=True)

            with d2:
                st.button("Clear", key="clear_text", on_click=lambda: st.session_state.update({"reset_triggered": True}))

# === ROUTING ===
page = st.experimental_get_query_params().get("page", ["Home"])[0]
if   page=="Home":            home_page()
elif page=="Manage EQ List":  create_eq_page()
elif page=="Search":          search_page()
elif page=="FAQ":             faq_page()
elif page=="Translation":     translation_page()
else: st.error(f"Unknown page: {page}")
